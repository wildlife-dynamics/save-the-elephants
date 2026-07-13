"""Tests for ecoscope_workflows_ext_ste.tasks.reporting._general_context.

Covers the private helpers (`build_report_context`, `prepare_general_context`,
`build_panel_rows`) and the single `wt_registry.register()`-decorated entry
point `general_template_context` (`register()` is a no-op at call time, so it
behaves as an ordinary function here).
"""

from __future__ import annotations

import os

import docx
import pytest
from docxtpl import DocxTemplate, InlineImage

from ecoscope_workflows_ext_ste.tasks.reporting._general_context import (
    build_panel_rows,
    build_report_context,
    general_template_context,
    prepare_general_context,
)


class TestBuildReportContext:
    def test_recognized_filenames_map_to_context_keys(self, tmp_path, make_png):
        make_png(name="collared_points.png")
        make_png(name="movement_tracks.png")

        ctx = build_report_context(str(tmp_path))

        assert ctx["collared_elephants_map"] == str(tmp_path / "collared_points.png")
        assert ctx["historical_current_tracks_map"] == str(tmp_path / "movement_tracks.png")
        # unset keys from filename_map default to None
        assert ctx["home_range_metrics_map"] is None

    def test_unrecognized_non_seasonal_png_is_ignored(self, tmp_path, make_png):
        make_png(name="readme.png")

        ctx = build_report_context(str(tmp_path))

        assert ctx["seasonal_images"] == []
        assert all(v is None for k, v in ctx.items() if k != "seasonal_images")

    def test_seasonal_pattern_extracts_subject_name(self, tmp_path, make_png):
        # pattern: ^[0-9a-f]+_(.+)\.png$
        make_png(name="abc123_elephant_seven.png")

        ctx = build_report_context(str(tmp_path))

        assert len(ctx["seasonal_images"]) == 1
        entry = ctx["seasonal_images"][0]
        assert entry["name"] == "elephant_seven"
        assert entry["image"] == str(tmp_path / "abc123_elephant_seven.png")

    def test_non_png_files_are_ignored(self, tmp_path, make_png):
        make_png(name="collared_points.jpg")  # wrong extension, glob only matches *.png

        ctx = build_report_context(str(tmp_path))

        assert ctx["collared_elephants_map"] is None

    def test_recurses_into_subdirectories(self, tmp_path, make_png):
        make_png(name="nested/collared_points.png")

        ctx = build_report_context(str(tmp_path))

        assert ctx["collared_elephants_map"] == str(tmp_path / "nested" / "collared_points.png")

    def test_empty_directory_returns_all_none_and_empty_seasonal(self, tmp_path):
        ctx = build_report_context(str(tmp_path))
        assert ctx["seasonal_images"] == []
        assert all(v is None for k, v in ctx.items() if k != "seasonal_images")


class TestPrepareGeneralContext:
    def _template(self, tmp_path):
        p = tmp_path / "t.docx"
        docx.Document().save(p)
        return DocxTemplate(str(p))

    def test_none_passes_through(self, tmp_path):
        tpl = self._template(tmp_path)
        assert prepare_general_context(None, tpl) is None

    def test_dict_recursion(self, tmp_path):
        tpl = self._template(tmp_path)
        result = prepare_general_context({"a": {"b": 1.9}}, tpl)
        assert result == {"a": {"b": 1}}

    def test_list_recursion(self, tmp_path):
        tpl = self._template(tmp_path)
        result = prepare_general_context([1.9, None, "text"], tpl)
        assert result == [1, None, "text"]

    def test_nonexistent_path_string_passes_through_unchanged(self, tmp_path):
        tpl = self._template(tmp_path)
        result = prepare_general_context("not_a_real_path.png", tpl)
        assert result == "not_a_real_path.png"

    def test_existing_non_image_path_passes_through_unchanged(self, tmp_path):
        tpl = self._template(tmp_path)
        text_file = tmp_path / "notes.txt"
        text_file.write_text("hi")
        result = prepare_general_context(str(text_file), tpl)
        assert result == str(text_file)

    def test_existing_image_path_becomes_inline_image(self, tmp_path, make_png):
        tpl = self._template(tmp_path)
        img = make_png(size=(200, 100))
        result = prepare_general_context(str(img), tpl, height=2.0, width=3.0)
        assert isinstance(result, InlineImage)

    def test_int_and_float_are_truncated_to_int(self, tmp_path):
        tpl = self._template(tmp_path)
        assert prepare_general_context(3.9, tpl) == 3
        assert prepare_general_context(-3.9, tpl) == -3
        assert prepare_general_context(4, tpl) == 4

    def test_bool_is_silently_coerced_via_int_branch(self, tmp_path):
        # NOTE: unlike build_docx_context in _mapbook_context.py (which special-cases
        # bool -> "Yes"/"No" *before* the int/float check), this transform() has no
        # bool branch. Since `bool` is a subclass of `int` in Python, True/False fall
        # into the `isinstance(value, (int, float))` branch and become 1/0. This is a
        # behavioral inconsistency between the two "reporting" context builders worth
        # flagging, not a crash -- documenting the actual (perhaps unintended) behavior.
        tpl = self._template(tmp_path)
        assert prepare_general_context(True, tpl) == 1
        assert prepare_general_context(False, tpl) == 0

    def test_other_types_pass_through(self, tmp_path):
        tpl = self._template(tmp_path)
        obj = object()
        assert prepare_general_context(obj, tpl) is obj


class TestBuildPanelRows:
    def test_even_number_pairs_up(self):
        images = [{"name": "a", "image": "a.png"}, {"name": "b", "image": "b.png"}]
        rows = build_panel_rows(images, cols=2)
        assert rows == [images]

    def test_odd_number_pads_last_row(self):
        images = [{"name": "a", "image": "a.png"}, {"name": "b", "image": "b.png"}, {"name": "c", "image": "c.png"}]
        rows = build_panel_rows(images, cols=2)
        assert len(rows) == 2
        assert rows[0] == images[:2]
        assert rows[1] == [images[2], {"name": "", "image": ""}]

    def test_empty_list_returns_empty_rows(self):
        assert build_panel_rows([], cols=2) == []

    def test_custom_column_count(self):
        images = [{"name": str(i)} for i in range(7)]
        rows = build_panel_rows(images, cols=3)
        assert len(rows) == 3
        assert len(rows[0]) == 3
        assert len(rows[-1]) == 3  # padded


class TestGeneralTemplateContext:
    def test_raises_if_template_path_missing(self, tmp_path):
        with pytest.raises(FileNotFoundError, match="Template file not found"):
            general_template_context(
                output_dir=str(tmp_path),
                template_path=str(tmp_path / "nope.docx"),
            )

    def test_raises_on_empty_template_path(self, tmp_path):
        with pytest.raises(ValueError, match="template_path is empty"):
            general_template_context(output_dir=str(tmp_path), template_path="  ")

    def test_raises_on_empty_output_dir(self, tmp_path, make_docx_template):
        template = make_docx_template(["x"])
        with pytest.raises(ValueError, match="output_dir is empty"):
            general_template_context(output_dir="   ", template_path=str(template))

    def test_raises_value_error_on_unloadable_template(self, tmp_path):
        garbage = tmp_path / "garbage.docx"
        garbage.write_bytes(b"not actually a docx")
        with pytest.raises(ValueError, match="Failed to load template"):
            general_template_context(output_dir=str(tmp_path / "out"), template_path=str(garbage))

    def test_happy_path_renders_and_saves(self, tmp_path, make_png):
        output_dir = tmp_path / "out"
        output_dir.mkdir()
        make_png(name="out/collared_points.png", size=(200, 100))

        template = tmp_path / "t.docx"
        doc = docx.Document()
        doc.add_paragraph("Map: {{ collared_elephants_map }}")
        doc.save(template)

        result_path = general_template_context(
            output_dir=str(output_dir),
            template_path=str(template),
        )

        assert result_path == str(output_dir / "general_context.docx")
        assert os.path.exists(result_path)

    def test_custom_filename(self, tmp_path):
        output_dir = tmp_path / "out"
        output_dir.mkdir()
        template = tmp_path / "t.docx"
        docx.Document().save(template)

        result_path = general_template_context(
            output_dir=str(output_dir),
            template_path=str(template),
            filename="custom.docx",
        )
        assert result_path == str(output_dir / "custom.docx")

    def test_file_scheme_paths_normalized(self, tmp_path):
        output_dir = tmp_path / "out"
        output_dir.mkdir()
        template = tmp_path / "t.docx"
        docx.Document().save(template)

        result_path = general_template_context(
            output_dir="file://" + str(output_dir),
            template_path="file://" + str(template),
        )
        assert not result_path.startswith("file://")
        assert os.path.exists(result_path)

    def test_render_failure_raises_value_error(self, tmp_path):
        # Malformed Jinja block (unclosed {% for %}) triggers a
        # jinja2.TemplateSyntaxError at render() time, which the function
        # wraps into a ValueError.
        output_dir = tmp_path / "out"
        output_dir.mkdir()
        template = tmp_path / "bad.docx"
        doc = docx.Document()
        doc.add_paragraph("{% for x in seasonal_panel %}{{ x }}")  # no endfor
        doc.save(template)

        with pytest.raises(ValueError, match="Failed to render or save document"):
            general_template_context(output_dir=str(output_dir), template_path=str(template))
