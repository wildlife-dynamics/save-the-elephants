"""Tests for ecoscope_workflows_ext_ste.tasks.reporting._mapbook_context.

This is the largest/most complex module in the `reporting` subpackage. It
covers private helpers (`_unwrap_skip`, `_format_area`, `_format_period`,
`_find_map_file`, `_fit_within_box`, `_make_inline_image`, `build_docx_context`,
`_default_filename`) plus the two `wt_registry.register()`-decorated entry
points (`register()` is a no-op at call time, so these behave as ordinary
functions here): `create_mapbook_context` and `render_mapbook_page`.
"""

from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from types import SimpleNamespace

import docx
import pandas as pd
import pytest
from docxtpl import DocxTemplate, InlineImage
from wt_task.skip import SKIP_SENTINEL

from ecoscope.platform.tasks.transformation._unit import Quantity, Unit
from ecoscope_workflows_ext_ste.tasks.reporting._mapbook_context import (
    MAP_SUFFIXES,
    TIME_FMT,
    _default_filename,
    _fit_within_box,
    _find_map_file,
    _format_area,
    _format_period,
    _make_inline_image,
    _unwrap_skip,
    build_docx_context,
    create_mapbook_context,
    render_mapbook_page,
)


# --------------------------------------------------------------------------
# _unwrap_skip
# --------------------------------------------------------------------------
class TestUnwrapSkip:
    def test_none_stays_none(self):
        assert _unwrap_skip(None) is None

    def test_skip_sentinel_becomes_none(self):
        assert _unwrap_skip(SKIP_SENTINEL) is None

    def test_plain_scalar_passes_through(self):
        assert _unwrap_skip("hello") == "hello"
        assert _unwrap_skip(42) == 42

    def test_list_with_sentinel_and_one_value_collapses_to_scalar(self):
        # items filtered to [5] -> single-item lists are unwrapped to a scalar,
        # not returned as a one-element list.
        assert _unwrap_skip([SKIP_SENTINEL, 5]) == 5

    def test_list_with_all_sentinels_becomes_none(self):
        assert _unwrap_skip([SKIP_SENTINEL, SKIP_SENTINEL]) is None

    def test_tuple_with_multiple_real_values_preserves_type_and_order(self):
        result = _unwrap_skip((1, 2))
        assert result == (1, 2)
        assert isinstance(result, tuple)

    def test_list_with_multiple_real_values_preserves_type(self):
        result = _unwrap_skip([1, 2, 3])
        assert result == [1, 2, 3]
        assert isinstance(result, list)

    def test_nested_list_is_recursively_unwrapped(self):
        # inner [SKIP_SENTINEL, 3] -> collapses to 3; outer -> [3, 4]
        result = _unwrap_skip([[SKIP_SENTINEL, 3], 4])
        assert result == [3, 4]

    def test_empty_list_becomes_none(self):
        assert _unwrap_skip([]) is None


# --------------------------------------------------------------------------
# _format_area / _format_period
# --------------------------------------------------------------------------
class TestFormatArea:
    def test_none_returns_na(self):
        assert _format_area(None) == "N/A"

    def test_object_with_none_value_returns_na(self):
        assert _format_area(SimpleNamespace(value=None, unit="km²")) == "N/A"

    def test_quantity_formats_value_and_unit(self):
        q = Quantity(value=12.34, unit=Unit.SQUARE_KILOMETER)
        assert _format_area(q) == "12.3 km²"

    def test_zero_value_is_not_treated_as_missing(self):
        q = Quantity(value=0, unit=Unit.SQUARE_METER)
        assert _format_area(q) == "0.0 m²"


class TestFormatPeriod:
    def test_none_returns_na(self):
        assert _format_period(None) == "N/A"

    def test_time_range_formats_since_and_until(self, make_time_range):
        tr = make_time_range(since="2024-06-01T08:30:00", until="2024-06-15T17:45:00")
        expected = f"{tr.since.strftime(TIME_FMT)} to {tr.until.strftime(TIME_FMT)}"
        assert _format_period(tr) == expected
        assert "Jun 2024" in _format_period(tr)


# --------------------------------------------------------------------------
# _find_map_file
# --------------------------------------------------------------------------
class TestFindMapFile:
    def test_exact_match_found(self, tmp_path):
        (tmp_path / "elephant_movement_tracks.png").write_bytes(b"x")
        result = _find_map_file(tmp_path, "elephant", "_movement_tracks")
        assert result == str(tmp_path / "elephant_movement_tracks.png")

    def test_jpg_extension_also_found(self, tmp_path):
        (tmp_path / "elephant_speedmap.jpg").write_bytes(b"x")
        result = _find_map_file(tmp_path, "elephant", "_speedmap")
        assert result == str(tmp_path / "elephant_speedmap.jpg")

    def test_glob_fallback_when_no_exact_match(self, tmp_path):
        (tmp_path / "elephant_homerange_v2.png").write_bytes(b"x")
        result = _find_map_file(tmp_path, "elephant", "_homerange")
        assert result == str(tmp_path / "elephant_homerange_v2.png")

    def test_multiple_glob_candidates_returns_most_recently_modified(self, tmp_path):
        older = tmp_path / "elephant_homerange_a.png"
        newer = tmp_path / "elephant_homerange_b.png"
        older.write_bytes(b"x")
        newer.write_bytes(b"x")
        now = datetime.now().timestamp()
        os.utime(older, (now - 100, now - 100))
        os.utime(newer, (now, now))

        result = _find_map_file(tmp_path, "elephant", "_homerange")
        assert result == str(newer)

    def test_no_match_returns_none(self, tmp_path):
        assert _find_map_file(tmp_path, "elephant", "_homerange") is None


# --------------------------------------------------------------------------
# create_mapbook_context
# --------------------------------------------------------------------------
class TestCreateMapbookContext:
    def test_df_none_uses_all_grouper_and_na_maps(self, tmp_path):
        ctx = create_mapbook_context(
            df=None,
            current_period=None,
            previous_period=None,
            period=None,
            grid_area=None,
            mcp_area=None,
            root_path=str(tmp_path),
        )
        assert ctx["grouper_value"] == "All"
        assert ctx["time_period"] == "N/A"
        assert ctx["previous_time_range"] == "N/A"
        assert ctx["period"] == "N/A"
        assert ctx["grid_area"] == "N/A"
        assert ctx["mcp_area"] == "N/A"
        for key in MAP_SUFFIXES:
            assert ctx[key] == "N/A"

    def test_skip_sentinel_inputs_treated_like_none(self, tmp_path):
        ctx = create_mapbook_context(
            df=SKIP_SENTINEL,
            current_period=SKIP_SENTINEL,
            previous_period=SKIP_SENTINEL,
            period=SKIP_SENTINEL,
            grid_area=SKIP_SENTINEL,
            mcp_area=SKIP_SENTINEL,
            root_path=str(tmp_path),
        )
        assert ctx["grouper_value"] == "All"
        assert ctx["period"] == "N/A"

    def test_single_subject_resolves_map_files(self, tmp_path):
        df = pd.DataFrame({"subject_name": ["Elephant One", "Elephant One"]})
        (tmp_path / "elephant_one_movement_tracks.png").write_bytes(b"x")

        ctx = create_mapbook_context(
            df=df,
            current_period=None,
            previous_period=None,
            period=None,
            grid_area=None,
            mcp_area=None,
            root_path=str(tmp_path),
        )
        assert ctx["grouper_value"] == "Elephant One"
        assert ctx["movement_tracks_map"] == str(tmp_path / "elephant_one_movement_tracks.png")
        # other suffixes weren't created on disk
        assert ctx["home_range_map"] == "N/A"

    def test_multiple_subject_names_falls_back_to_all(self, tmp_path):
        df = pd.DataFrame({"subject_name": ["A", "B"]})
        (tmp_path / "a_movement_tracks.png").write_bytes(b"x")

        ctx = create_mapbook_context(
            df=df,
            current_period=None,
            previous_period=None,
            period=None,
            grid_area=None,
            mcp_area=None,
            root_path=str(tmp_path),
        )
        # safe_name stays None when >1 unique subject_name, so no per-subject
        # file lookup happens even though a plausible file exists on disk.
        assert ctx["grouper_value"] == "All"
        assert ctx["movement_tracks_map"] == "N/A"

    def test_missing_subject_name_column_falls_back_to_all(self, tmp_path):
        df = pd.DataFrame({"other_col": [1, 2, 3]})
        ctx = create_mapbook_context(
            df=df,
            current_period=None,
            previous_period=None,
            period=None,
            grid_area=None,
            mcp_area=None,
            root_path=str(tmp_path),
        )
        assert ctx["grouper_value"] == "All"

    def test_period_formatting_and_skip_unwrap(self, tmp_path):
        ctx = create_mapbook_context(
            df=None,
            current_period=None,
            previous_period=None,
            period=[SKIP_SENTINEL, 5.5],
            grid_area=None,
            mcp_area=None,
            root_path=str(tmp_path),
        )
        assert ctx["period"] == "5.50"

    def test_time_periods_and_areas_formatted(self, tmp_path, make_time_range):
        current = make_time_range(since="2024-01-01T00:00:00", until="2024-01-31T00:00:00")
        previous = make_time_range(since="2023-12-01T00:00:00", until="2023-12-31T00:00:00")
        grid = Quantity(value=10.0, unit=Unit.SQUARE_KILOMETER)
        mcp = Quantity(value=20.0, unit=Unit.SQUARE_KILOMETER)

        ctx = create_mapbook_context(
            df=None,
            current_period=current,
            previous_period=previous,
            period=3.0,
            grid_area=grid,
            mcp_area=mcp,
            root_path=str(tmp_path),
        )
        assert ctx["time_period"] == _format_period(current)
        assert ctx["previous_time_range"] == _format_period(previous)
        assert ctx["grid_area"] == "10.0 km²"
        assert ctx["mcp_area"] == "20.0 km²"
        assert ctx["period"] == "3.00"

    def test_root_path_with_file_scheme_is_normalized(self, tmp_path):
        df = pd.DataFrame({"subject_name": ["Foo"]})
        (tmp_path / "foo_homerange.png").write_bytes(b"x")

        ctx = create_mapbook_context(
            df=df,
            current_period=None,
            previous_period=None,
            period=None,
            grid_area=None,
            mcp_area=None,
            root_path="file://" + str(tmp_path),
        )
        assert ctx["home_range_map"] == str(tmp_path / "foo_homerange.png")


# --------------------------------------------------------------------------
# _fit_within_box / _make_inline_image
# --------------------------------------------------------------------------
class TestFitWithinBox:
    def test_wide_image_constrained_by_width(self, make_png):
        path = make_png(size=(1000, 500))
        w, h = _fit_within_box(str(path), box_w_cm=10, box_h_cm=10)
        # scale = min(10/1000, 10/500) = min(0.01, 0.02) = 0.01
        assert w == pytest.approx(10.0)
        assert h == pytest.approx(5.0)

    def test_tall_image_constrained_by_height(self, make_png):
        path = make_png(size=(500, 1000))
        w, h = _fit_within_box(str(path), box_w_cm=10, box_h_cm=10)
        assert w == pytest.approx(5.0)
        assert h == pytest.approx(10.0)


class TestMakeInlineImage:
    def _template(self, tmp_path):
        p = tmp_path / "t.docx"
        docx.Document().save(p)
        return DocxTemplate(str(p))

    def test_valid_image_returns_inline_image(self, tmp_path, make_png):
        tpl = self._template(tmp_path)
        img_path = make_png(size=(200, 100))
        result = _make_inline_image(tpl, str(img_path), box_w_cm=5, box_h_cm=5)
        assert isinstance(result, InlineImage)

    def test_corrupt_image_returns_none_not_raise(self, tmp_path):
        tpl = self._template(tmp_path)
        corrupt = tmp_path / "corrupt.png"
        corrupt.write_bytes(b"not a real image")
        result = _make_inline_image(tpl, str(corrupt), box_w_cm=5, box_h_cm=5)
        assert result is None


# --------------------------------------------------------------------------
# build_docx_context
# --------------------------------------------------------------------------
class TestBuildDocxContext:
    def _template(self, tmp_path):
        p = tmp_path / "t.docx"
        docx.Document().save(p)
        return DocxTemplate(str(p))

    def test_none_becomes_empty_string(self, tmp_path):
        tpl = self._template(tmp_path)
        result = build_docx_context({"a": None}, tpl)
        assert result == {"a": ""}

    def test_bool_becomes_yes_no(self, tmp_path):
        tpl = self._template(tmp_path)
        result = build_docx_context({"t": True, "f": False}, tpl)
        assert result == {"t": "Yes", "f": "No"}

    def test_int_gets_comma_formatted(self, tmp_path):
        tpl = self._template(tmp_path)
        result = build_docx_context({"n": 1234567}, tpl)
        assert result == {"n": "1,234,567"}

    def test_float_gets_comma_and_two_decimals(self, tmp_path):
        tpl = self._template(tmp_path)
        result = build_docx_context({"n": 1234.5}, tpl)
        assert result == {"n": "1,234.50"}

    def test_existing_image_path_becomes_inline_image(self, tmp_path, make_png):
        tpl = self._template(tmp_path)
        img = make_png(size=(200, 100))
        result = build_docx_context({"pic": str(img)}, tpl)
        assert isinstance(result["pic"], InlineImage)

    def test_nonexistent_image_path_string_passes_through_unchanged(self, tmp_path):
        tpl = self._template(tmp_path)
        missing = str(tmp_path / "missing.png")
        result = build_docx_context({"pic": missing}, tpl)
        assert result["pic"] == missing

    def test_existing_non_image_suffix_path_passes_through_unchanged(self, tmp_path):
        tpl = self._template(tmp_path)
        text_file = tmp_path / "notes.txt"
        text_file.write_text("hi")
        result = build_docx_context({"note": str(text_file)}, tpl)
        assert result["note"] == str(text_file)

    def test_plain_string_passes_through(self, tmp_path):
        tpl = self._template(tmp_path)
        result = build_docx_context({"title": "Elephant Report"}, tpl)
        assert result["title"] == "Elephant Report"

    def test_other_types_pass_through_without_recursion(self, tmp_path):
        # Unlike _general_context.prepare_general_context, build_docx_context
        # does NOT recurse into dict/list values -- they pass through as-is.
        tpl = self._template(tmp_path)
        nested = {"x": 1}
        result = build_docx_context({"nested": nested}, tpl)
        assert result["nested"] is nested

    def test_no_keys_added_or_removed(self, tmp_path):
        tpl = self._template(tmp_path)
        context = {"a": None, "b": True, "c": 1, "d": 1.0, "e": "text"}
        result = build_docx_context(context, tpl)
        assert set(result.keys()) == set(context.keys())

    def test_custom_box_dimensions_used_for_scaling(self, tmp_path, make_png):
        tpl = self._template(tmp_path)
        img = make_png(size=(1000, 1000))
        result = build_docx_context({"pic": str(img)}, tpl, box_h_cm=2, box_w_cm=2)
        assert isinstance(result["pic"], InlineImage)


# --------------------------------------------------------------------------
# _default_filename
# --------------------------------------------------------------------------
class TestDefaultFilename:
    def test_uses_sanitized_grouper_value(self):
        assert _default_filename({"grouper_value": "Elephant One"}) == "elephant_one.docx"

    def test_underscore_already_safe_value_unchanged(self):
        assert _default_filename({"grouper_value": "elephant_02"}) == "elephant_02.docx"

    def test_strips_leading_trailing_whitespace(self):
        assert _default_filename({"grouper_value": "  Leading Space  "}) == "leading_space.docx"

    def test_na_grouper_value_falls_back_to_uuid(self):
        name = _default_filename({"grouper_value": "N/A"})
        assert re.match(r"^[0-9a-f]{8}\.docx$", name)

    def test_na_case_insensitive(self):
        name = _default_filename({"grouper_value": "n/a"})
        assert re.match(r"^[0-9a-f]{8}\.docx$", name)

    def test_missing_grouper_value_key_falls_back_to_uuid(self):
        name = _default_filename({})
        assert re.match(r"^[0-9a-f]{8}\.docx$", name)

    def test_empty_grouper_value_falls_back_to_uuid(self):
        name = _default_filename({"grouper_value": "   "})
        assert re.match(r"^[0-9a-f]{8}\.docx$", name)


# --------------------------------------------------------------------------
# render_mapbook_page
# --------------------------------------------------------------------------
class TestRenderMapbookPage:
    def test_raises_on_empty_template_path(self, tmp_path):
        with pytest.raises(ValueError, match="template_path is empty"):
            render_mapbook_page(template_path="  ", output_dir=str(tmp_path), context={})

    def test_raises_on_empty_output_dir(self, tmp_path, make_docx_template):
        template = make_docx_template(["x"])
        with pytest.raises(ValueError, match="output_dir is empty"):
            render_mapbook_page(template_path=str(template), output_dir="  ", context={})

    def test_raises_file_not_found_for_missing_template(self, tmp_path):
        with pytest.raises(FileNotFoundError, match="Template file not found"):
            render_mapbook_page(
                template_path=str(tmp_path / "nope.docx"),
                output_dir=str(tmp_path / "out"),
                context={},
            )

    def test_happy_path_default_filename_from_grouper(self, make_docx_template, tmp_path, read_docx_text):
        template = make_docx_template(["Subject: {{ grouper_value }}"])
        output_dir = tmp_path / "out"

        result = render_mapbook_page(
            template_path=str(template),
            output_dir=str(output_dir),
            context={"grouper_value": "Elephant One"},
        )
        assert result == str(output_dir / "elephant_one.docx")
        assert os.path.exists(result)
        assert "Subject: Elephant One" in read_docx_text(Path(result))

    def test_custom_filename_overrides_default(self, make_docx_template, tmp_path):
        template = make_docx_template(["x"])
        output_dir = tmp_path / "out"

        result = render_mapbook_page(
            template_path=str(template),
            output_dir=str(output_dir),
            context={"grouper_value": "Elephant One"},
            filename="custom_page.docx",
        )
        assert result == str(output_dir / "custom_page.docx")

    def test_missing_image_default_is_lenient_and_still_renders(self, make_docx_template, tmp_path, capsys):
        template = make_docx_template(["Map: {{ movement_tracks_map }}"])
        output_dir = tmp_path / "out"
        missing_path = str(tmp_path / "does_not_exist.png")

        result = render_mapbook_page(
            template_path=str(template),
            output_dir=str(output_dir),
            context={"movement_tracks_map": missing_path},
            strict_images=False,
        )
        assert os.path.exists(result)
        captured = capsys.readouterr()
        assert "image not found" in captured.out

    def test_missing_image_strict_mode_raises(self, make_docx_template, tmp_path):
        template = make_docx_template(["Map: {{ movement_tracks_map }}"])
        output_dir = tmp_path / "out"
        missing_path = str(tmp_path / "does_not_exist.png")

        with pytest.raises(FileNotFoundError, match="image not found"):
            render_mapbook_page(
                template_path=str(template),
                output_dir=str(output_dir),
                context={"movement_tracks_map": missing_path},
                strict_images=True,
            )

    def test_non_string_and_empty_context_values_are_skipped_in_validation(self, make_docx_template, tmp_path):
        template = make_docx_template(["Count: {{ count }}"])
        output_dir = tmp_path / "out"

        # None, 0, "", and non-path ints/floats should not trip the image
        # validation loop (it only inspects non-empty strings).
        result = render_mapbook_page(
            template_path=str(template),
            output_dir=str(output_dir),
            context={"count": 5, "empty": "", "none_val": None, "zero": 0},
            strict_images=True,
        )
        assert os.path.exists(result)

    def test_valid_image_embedded_successfully(self, make_docx_template, make_png, tmp_path):
        template = make_docx_template(["Map: {{ movement_tracks_map }}"])
        output_dir = tmp_path / "out"
        img = make_png(size=(400, 200))

        result = render_mapbook_page(
            template_path=str(template),
            output_dir=str(output_dir),
            context={"movement_tracks_map": str(img)},
            strict_images=True,
        )
        assert os.path.exists(result)

    def test_raises_value_error_on_unloadable_template(self, tmp_path):
        # Whether a corrupt docx fails at `DocxTemplate(template_path)` (caught
        # by the "Failed to load template" branch) or later at `.render()`/
        # `.save()` (caught by the "Failed to render or save" branch) depends
        # on the installed docxtpl/python-docx version, which can differ across
        # platforms/solve groups -- so only assert on the invariant part: the
        # underlying package error surfaces as a ValueError.
        garbage = tmp_path / "garbage.docx"
        garbage.write_bytes(b"not actually a docx")
        with pytest.raises(ValueError, match="Package not found"):
            render_mapbook_page(
                template_path=str(garbage),
                output_dir=str(tmp_path / "out"),
                context={},
            )

    def test_render_failure_raises_value_error(self, tmp_path):
        output_dir = tmp_path / "out"
        template = tmp_path / "bad.docx"
        doc = docx.Document()
        doc.add_paragraph("{% for x in items %}{{ x }}")  # unterminated block
        doc.save(template)

        with pytest.raises(ValueError, match="Failed to render or save"):
            render_mapbook_page(
                template_path=str(template),
                output_dir=str(output_dir),
                context={"items": [1, 2, 3]},
            )

    def test_file_scheme_paths_normalized(self, make_docx_template, tmp_path):
        template = make_docx_template(["x"])
        output_dir = tmp_path / "out"

        result = render_mapbook_page(
            template_path="file://" + str(template),
            output_dir="file://" + str(output_dir),
            context={},
        )
        assert not result.startswith("file://")
        assert os.path.exists(result)

    def test_custom_box_dimensions_passed_through(self, make_docx_template, make_png, tmp_path):
        template = make_docx_template(["Map: {{ pic }}"])
        output_dir = tmp_path / "out"
        img = make_png(size=(2000, 1000))

        result = render_mapbook_page(
            template_path=str(template),
            output_dir=str(output_dir),
            context={"pic": str(img)},
            box_h_cm=3.0,
            box_w_cm=3.0,
        )
        assert os.path.exists(result)
