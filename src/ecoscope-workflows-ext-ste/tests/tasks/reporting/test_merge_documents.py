"""Tests for ecoscope_workflows_ext_ste.tasks.reporting._merge_documents.

Covers the single `wt_registry.register()`-decorated function
(`register()` is a no-op at call time, so it's exercised as an ordinary
function here): `merge_docx_documents`, which combines a cover page with a
list of context-page .docx files into one document using docxcompose.
"""

from __future__ import annotations

import os
from pathlib import Path

import docx
import pytest
from wt_task.skip import SKIP_SENTINEL

from ecoscope_workflows_ext_ste.tasks.reporting._merge_documents import merge_docx_documents


def _make_doc(path: Path, text: str) -> Path:
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _paragraphs(path: str) -> list[str]:
    return [p.text for p in docx.Document(path).paragraphs]


class TestMergeDocxDocuments:
    def test_raises_if_cover_page_missing(self, tmp_path):
        with pytest.raises(FileNotFoundError, match="Cover page file not found"):
            merge_docx_documents(
                cover_page_path=str(tmp_path / "nope.docx"),
                context_page_items=[],
                output_dir=str(tmp_path / "out"),
            )

    def test_raises_on_empty_output_dir(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        with pytest.raises(ValueError, match="output_dir is empty"):
            merge_docx_documents(
                cover_page_path=str(cover),
                context_page_items=[],
                output_dir="   ",
            )

    def test_cover_only_when_no_context_pages(self, tmp_path, capsys):
        cover = _make_doc(tmp_path / "cover.docx", "COVER PAGE")
        output_dir = tmp_path / "out"

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[],
            output_dir=str(output_dir),
        )
        assert os.path.exists(result)
        assert _paragraphs(result) == ["COVER PAGE"]
        assert "No valid context pages to merge" in capsys.readouterr().out

    def test_merges_cover_and_context_pages_in_name_order(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        feb = _make_doc(tmp_path / "feb_page.docx", "FEB CONTEXT")
        jan = _make_doc(tmp_path / "jan_page.docx", "JAN CONTEXT")

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[str(feb), str(jan)],
            output_dir=str(tmp_path / "out"),
            order_by="name",
        )
        paragraphs = _paragraphs(result)
        assert paragraphs == ["COVER", "FEB CONTEXT", "JAN CONTEXT"]  # alphabetical: feb < jan

    def test_default_filename_has_timestamp_pattern(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[],
            output_dir=str(tmp_path / "out"),
        )
        assert Path(result).name.startswith("overall_report_")
        assert Path(result).suffix == ".docx"

    def test_custom_filename_without_extension_gets_docx_appended(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[],
            output_dir=str(tmp_path / "out"),
            filename="my_report",
        )
        assert Path(result).name == "my_report.docx"

    def test_custom_filename_with_extension_is_untouched(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[],
            output_dir=str(tmp_path / "out"),
            filename="MyReport.DOCX",
        )
        assert Path(result).name == "MyReport.DOCX"

    def test_skip_sentinel_items_are_skipped(self, tmp_path, capsys):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        jan = _make_doc(tmp_path / "jan_page.docx", "JAN CONTEXT")

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[SKIP_SENTINEL, str(jan)],
            output_dir=str(tmp_path / "out"),
        )
        assert _paragraphs(result) == ["COVER", "JAN CONTEXT"]
        assert "Skipping page item 0" in capsys.readouterr().out

    def test_none_items_are_skipped(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        jan = _make_doc(tmp_path / "jan_page.docx", "JAN CONTEXT")

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[None, str(jan)],
            output_dir=str(tmp_path / "out"),
        )
        assert _paragraphs(result) == ["COVER", "JAN CONTEXT"]

    def test_nonexistent_path_items_are_skipped_with_message(self, tmp_path, capsys):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[str(tmp_path / "missing.docx")],
            output_dir=str(tmp_path / "out"),
        )
        assert _paragraphs(result) == ["COVER"]
        assert "file not found" in capsys.readouterr().out

    def test_list_item_extracts_first_existing_string_path(self, tmp_path):
        # extract_path: for list/tuple items, pick the first str element that
        # exists on disk (as long as none of the elements is SKIP_SENTINEL).
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        jan = _make_doc(tmp_path / "jan_page.docx", "JAN CONTEXT")

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[[str(tmp_path / "missing.docx"), str(jan)]],
            output_dir=str(tmp_path / "out"),
        )
        assert _paragraphs(result) == ["COVER", "JAN CONTEXT"]

    def test_list_item_containing_skip_sentinel_is_skipped_entirely(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        jan = _make_doc(tmp_path / "jan_page.docx", "JAN CONTEXT")

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[[SKIP_SENTINEL, str(jan)]],
            output_dir=str(tmp_path / "out"),
        )
        # Even though `jan` exists, the presence of SKIP_SENTINEL anywhere in
        # the tuple/list short-circuits extract_path to None for that item.
        assert _paragraphs(result) == ["COVER"]

    def test_order_by_month_sorts_pages_by_calendar_month(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        march = _make_doc(tmp_path / "march_report.docx", "MARCH")
        jan = _make_doc(tmp_path / "january_report.docx", "JANUARY")
        feb = _make_doc(tmp_path / "feb_report.docx", "FEBRUARY")

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[str(march), str(jan), str(feb)],
            output_dir=str(tmp_path / "out"),
            order_by="month",
        )
        assert _paragraphs(result) == ["COVER", "JANUARY", "FEBRUARY", "MARCH"]

    def test_order_by_month_unmatched_filenames_sort_after_matched_ones(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        jan = _make_doc(tmp_path / "january_report.docx", "JANUARY")
        unrelated = _make_doc(tmp_path / "summary.docx", "SUMMARY")

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[str(unrelated), str(jan)],
            output_dir=str(tmp_path / "out"),
            order_by="month",
        )
        assert _paragraphs(result) == ["COVER", "JANUARY", "SUMMARY"]

    def test_order_by_input_preserves_original_order(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        z_page = _make_doc(tmp_path / "z_page.docx", "Z CONTENT")
        a_page = _make_doc(tmp_path / "a_page.docx", "A CONTENT")

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[str(z_page), str(a_page)],
            output_dir=str(tmp_path / "out"),
            order_by="input",
        )
        # order_by="input" doesn't sort at all -- items retain the order they
        # were validated/appended in (input order).
        assert _paragraphs(result) == ["COVER", "Z CONTENT", "A CONTENT"]

    def test_output_dir_created_if_missing(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        output_dir = tmp_path / "brand" / "new" / "dir"
        assert not output_dir.exists()

        result = merge_docx_documents(
            cover_page_path=str(cover),
            context_page_items=[],
            output_dir=str(output_dir),
        )
        assert output_dir.exists()
        assert os.path.exists(result)

    def test_file_scheme_paths_normalized(self, tmp_path):
        cover = _make_doc(tmp_path / "cover.docx", "COVER")
        result = merge_docx_documents(
            cover_page_path="file://" + str(cover),
            context_page_items=[],
            output_dir="file://" + str(tmp_path / "out"),
        )
        assert not result.startswith("file://")
        assert os.path.exists(result)
