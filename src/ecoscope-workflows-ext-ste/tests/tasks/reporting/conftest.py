"""Shared fixtures for ecoscope_workflows_ext_ste.tasks.reporting tests.

The functions under test are registered via `wt_registry.register()`, which is
a no-op decorator at call time, so every registered function below is called
directly as plain Python -- no workflow engine involved.

Fixtures here focus on building minimal, real inputs (docx templates via
python-docx, images via Pillow, TimeRange/Quantity model instances) rather
than mocking the document-generation libraries themselves, since docxtpl /
docxcompose / Pillow are all installed in this environment and cheap to run
for real.
"""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import docx
import pytest
from PIL import Image

from ecoscope.platform.tasks.filter._filter import TimeRange, TimezoneInfo
from ecoscope.platform.tasks.transformation._unit import Quantity, Unit


@pytest.fixture
def utc_timezone() -> TimezoneInfo:
    return TimezoneInfo(label="UTC", tzCode="UTC", name="UTC", utc_offset="+00:00")


@pytest.fixture
def make_time_range(utc_timezone) -> Callable[..., TimeRange]:
    """Factory for TimeRange instances with sane defaults (both tz-naive, UTC)."""

    def _make(since="2024-01-01T00:00:00", until="2024-02-01T00:00:00", **kwargs) -> TimeRange:
        kwargs.setdefault("timezone", utc_timezone)
        return TimeRange(since=since, until=until, **kwargs)

    return _make


@pytest.fixture
def make_quantity() -> Callable[..., Quantity]:
    def _make(value=12.34, unit=Unit.SQUARE_KILOMETER) -> Quantity:
        return Quantity(value=value, unit=unit)

    return _make


@pytest.fixture
def make_png(tmp_path) -> Callable[..., Path]:
    """Factory that writes a small PNG of a given pixel size (and optional dpi)
    under tmp_path, returning its path as a string-friendly Path.
    """

    counter = {"n": 0}

    def _make(
        name: str | None = None,
        size: tuple[int, int] = (192, 96),
        dpi: tuple[int, int] | None = None,
        color=(255, 0, 0),
    ) -> Path:
        counter["n"] += 1
        filename = name or f"image_{counter['n']}.png"
        path = tmp_path / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        img = Image.new("RGB", size, color=color)
        if dpi is not None:
            img.save(path, dpi=dpi)
        else:
            img.save(path)
        return path

    return _make


@pytest.fixture
def make_docx_template(tmp_path) -> Callable[..., Path]:
    """Factory that writes a minimal .docx with the given paragraph strings
    (each may contain Jinja/docxtpl placeholders like '{{ field }}') and
    returns its path.
    """

    counter = {"n": 0}

    def _make(paragraphs: list[str], name: str | None = None) -> Path:
        counter["n"] += 1
        filename = name or f"template_{counter['n']}.docx"
        path = tmp_path / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = docx.Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    return _make


@pytest.fixture
def read_docx_text() -> Callable[[Path], list[str]]:
    """Return the paragraph texts of a saved .docx, for asserting render output."""

    def _read(path: Path) -> list[str]:
        d = docx.Document(str(path))
        return [p.text for p in d.paragraphs]

    return _read
