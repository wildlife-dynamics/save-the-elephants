import os
import uuid
import logging
from pathlib import Path
from pydantic import Field
from docx.shared import Cm
from datetime import datetime
from docxtpl import DocxTemplate, InlineImage
from ecoscope_workflows_core.decorators import task
from typing import Annotated, Optional, Dict, Any, Union
from ecoscope_workflows_core.tasks.filter._filter import TimeRange
from ecoscope_workflows_core.tasks.transformation._unit import Quantity
from ecoscope_workflows_ext_custom.tasks.io._path_utils import remove_file_scheme
from ecoscope_workflows_core.annotations import AnyDataFrame

logger = logging.getLogger(__name__)


@task
def create_context_page(
    template_path: Annotated[
        str,
        Field(
            description="Path to the .docx template file.",
        ),
    ],
    output_dir: Annotated[
        str,
        Field(
            description="Directory to save the generated .docx file.",
        ),
    ],
    context: Annotated[
        dict,
        Field(
            description="Dictionary with context values for the template.",
        ),
    ],
    logo_width_cm: Annotated[
        float,
        Field(
            description="Width of the logo in centimeters.",
        ),
    ] = 7.7,
    logo_height_cm: Annotated[
        float,
        Field(
            description="Height of the logo in centimeters.",
        ),
    ] = 1.93,
    filename: Annotated[
        Optional[str],
        Field(
            description="Optional filename . If not provided, a random UUID-based filename will be generated.",
            exclude=True,
        ),
    ] = None,
) -> Annotated[
    str,
    Field(
        description="Full path to the generated .docx file.",
    ),
]:
    """
    Create a context page document from a template and context dictionary.

    Args:
        template_path (str): Path to the .docx template file.
        output_dir (str): Directory to save the generated .docx file.
        context (dict): Dictionary with context values for the template.
        logo_width_cm (float): Width of the logo in centimeters. Default is 7.7.
        logo_height_cm (float): Height of the logo in centimeters. Default is 1.93.
        filename (str, optional): Optional filename for the generated file.
            If not provided, a random UUID-based filename will be generated.

    Returns:
        str: Full path to the generated .docx file.
    """
    logger.info("Starting create_context_page task.")
    # Normalize paths
    template_path = remove_file_scheme(template_path)
    output_dir = remove_file_scheme(output_dir)

    # Validate paths
    if not template_path.strip():
        raise ValueError("template_path is empty after normalization")
    if not output_dir.strip():
        raise ValueError("output_dir is empty after normalization")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        filename = "context_page_.docx"
    output_path = Path(output_dir) / filename

    doc = DocxTemplate(template_path)
    if "org_logo_path" in context and os.path.exists(context["org_logo_path"]):
        context["org_logo"] = InlineImage(
            doc,
            context["org_logo_path"],
            width=Cm(logo_width_cm),
            height=Cm(logo_height_cm),
        )
    doc.render(context)
    doc.save(output_path)
    logger.info(f"Context page document created at: {output_path}")
    return str(output_path)


@task
def create_mapbook_ctx_cover(
    count: int,
    org_logo_path: Union[str, Path, None],
    report_period: TimeRange,
    prepared_by: str,
) -> Dict[str, Optional[str]]:
    """
    Build a dictionary with the mapbook report template values.

    Args:
        count (int): Total number of subjects or records.
        org_logo_path (Union[str, Path, None]): Path to the org logo, or None if unavailable.
        report_period (TimeRange): Object with 'since', 'until', and 'time_format' attributes.
        prepared_by (str): Name of the person or organization preparing the report.

    Returns:
        Dict[str, Optional[str]]: Structured dictionary with formatted metadata.
    """
    # Resolve logo path — guard against None before calling remove_file_scheme
    resolved_logo_path: Optional[str] = None
    if org_logo_path is not None:
        resolved_logo_path = remove_file_scheme(org_logo_path)
        if not resolved_logo_path.strip():
            raise ValueError("org_logo_path is empty after normalization")

    formatted_date = datetime.now()
    formatted_date_str = formatted_date.strftime("%Y-%m-%d %H:%M:%S")
    fmt = getattr(report_period, "time_format", "%Y-%m-%d")
    formatted_time_range = f"{report_period.since.strftime(fmt)} to {report_period.until.strftime(fmt)}"

    return {
        "report_id": f"REP-{uuid.uuid4().hex[:8].upper()}",
        "subject_count": str(count),
        "org_logo_path": resolved_logo_path,  # None if no logo was provided
        "time_generated": formatted_date_str,
        "report_period": formatted_time_range,
        "prepared_by": prepared_by,
    }


def validate_image_path(field_name: str, path: str) -> None:
    """Validate that an image file exists and has valid extension."""
    normalized_path = remove_file_scheme(path)

    if not os.path.exists(normalized_path):
        raise FileNotFoundError(f"Image file for '{field_name}' not found: {normalized_path}")

    valid_extensions = {".png", ".jpg", ".jpeg"}
    if Path(normalized_path).suffix.lower() not in valid_extensions:
        raise ValueError(
            f"Invalid image format for '{field_name}': {Path(normalized_path).suffix}. "
            f"Expected one of {valid_extensions}"
        )

    logging.info(f" Validated image for '{field_name}': {normalized_path}")


@task
def create_mapbook_grouper_ctx(
    current_period: TimeRange,
    previous_period: TimeRange,
    period: float | None,
    grouper_name: tuple | list | str | None,
    df: AnyDataFrame,
    grid_area: Quantity | None,
    mcp_area: Quantity | None,
    map_paths: list | None,
) -> Dict[str, Optional[str]]:
    TIME_FMT = "%d %b %Y %H:%M:%S"

    # Format current period
    current_period_str = None
    if current_period:
        current_period_str = (
            f"{current_period.since.strftime(TIME_FMT)} " f"to {current_period.until.strftime(TIME_FMT)}"
        )

    # Format previous period
    previous_period_str = None
    if previous_period:
        previous_period_str = (
            f"{previous_period.since.strftime(TIME_FMT)} " f"to {previous_period.until.strftime(TIME_FMT)}"
        )

    # Format period
    period_str = f"{period:.2f}" if isinstance(period, float) else str(period)

    # Format areas with None handling
    grid_area_str = "N/A"
    if grid_area and grid_area.value is not None:
        grid_area_str = f"{grid_area.value:.1f} {grid_area.unit}"

    mcp_area_str = "N/A"
    if mcp_area and mcp_area.value is not None:
        mcp_area_str = f"{mcp_area.value:.1f} {mcp_area.unit}"

    logging.info(f"grid area: {grid_area_str} || mcp area: {mcp_area_str}")

    # Extract grouper value dynamically
    grouper_value = "All"
    logging.info(f"grouper name raw: {grouper_name} (type: {type(grouper_name)})")

    if grouper_name:
        if isinstance(grouper_name, str):
            grouper_value = grouper_name
        elif isinstance(grouper_name, (list, tuple)) and len(grouper_name) > 0:
            # Check if it's a tuple structure like (('index_name', 'All'),)
            first_item = grouper_name[0]

            # Handle tuple structure (('index_name', 'All'),)
            if isinstance(first_item, tuple) and len(first_item) == 2:
                key, value = first_item
                if key == "index_name" and value == "All":
                    grouper_value = "All"
                else:
                    grouper_value = str(value)
                logging.info(f"Extracted from tuple structure: {grouper_value}")
            # Handle grouper objects
            elif hasattr(first_item, "__class__"):
                grouper = first_item
                grouper_type = grouper.__class__.__name__
                logging.info(f"grouper_type: {grouper_type}")

                if grouper_type == "ValueGrouper":
                    index_name = getattr(grouper, "index_name", None)
                    if df is not None and index_name and index_name in df.columns:
                        unique_values = df[index_name].unique()
                        if len(unique_values) == 1:
                            grouper_value = str(unique_values[0])
                        else:
                            grouper_value = index_name
                    else:
                        grouper_value = index_name if index_name else "Value"

                elif grouper_type == "TemporalGrouper":
                    grouper_value = _format_temporal_grouper(grouper, df)

                elif grouper_type == "AllGrouper":
                    grouper_value = "All"
                else:
                    grouper_value = str(grouper_name)
            else:
                grouper_value = str(first_item)
        else:
            grouper_value = str(grouper_name)

    logging.info(f"grouper_value: {grouper_value}")

    # Map parsing with None handling
    map_suffixes = {
        "movement_tracks_map": "movement_tracks.png",
        "home_range_map": "homerange.png",
        "speed_map": "speedmap.png",
        "speed_raster_map": "mean_speed_raster.png",
        "night_day_ecomap": "day_night.png",
        "seasonal_map": "seasonal_home_range.png",
    }

    mapbook_png_paths = {key: None for key in map_suffixes.keys()}
    if map_paths:
        for path in map_paths:
            if path:
                for key, suffix in map_suffixes.items():
                    if path.endswith(suffix):
                        mapbook_png_paths[key] = path
                        break

    # Build context
    ctx = {
        "time_period": current_period_str,
        "previous_time_range": previous_period_str,
        "period": period_str,
        "grid_area": grid_area_str,
        "mcp_area": mcp_area_str,
        "grouper_value": grouper_value,
        **mapbook_png_paths,
    }

    logging.info(f"Context: {ctx}")
    return ctx


def _format_temporal_grouper(grouper: Any, df: AnyDataFrame) -> str:
    """
    Format temporal grouper with human-readable labels.
    Extracts actual temporal values from the segment_start column.
    """
    # Try to get the temporal_index attribute
    temporal_index = getattr(grouper, "temporal_index", None)

    if temporal_index and df is not None:
        # Get the directive (e.g., '%B' for full month name, '%A' for day name)
        directive = getattr(temporal_index, "directive", None)

        # Check if segment_start column exists
        if "segment_start" in df.columns:
            try:
                # Get unique dates from segment_start
                dates = df["segment_start"].dropna()

                if len(dates) > 0:
                    # Format based on directive
                    if directive == "%B":  # Full month name
                        unique_months = dates.dt.strftime("%B").unique()
                        if len(unique_months) == 1:
                            return unique_months[0]
                        elif len(unique_months) > 1:
                            return f"{unique_months[0]} - {unique_months[-1]}"

                    elif directive == "%b":  # Abbreviated month name
                        unique_months = dates.dt.strftime("%b").unique()
                        if len(unique_months) == 1:
                            return unique_months[0]
                        elif len(unique_months) > 1:
                            return f"{unique_months[0]} - {unique_months[-1]}"

                    elif directive == "%A":  # Full day name
                        unique_days = dates.dt.strftime("%A").unique()
                        if len(unique_days) == 1:
                            return unique_days[0]
                        elif len(unique_days) > 1:
                            return f"{unique_days[0]} - {unique_days[-1]}"

                    elif directive == "%a":  # Abbreviated day name
                        unique_days = dates.dt.strftime("%a").unique()
                        if len(unique_days) == 1:
                            return unique_days[0]
                        elif len(unique_days) > 1:
                            return f"{unique_days[0]} - {unique_days[-1]}"

                    elif directive == "%Y":  # Year
                        unique_years = dates.dt.strftime("%Y").unique()
                        if len(unique_years) == 1:
                            return unique_years[0]
                        elif len(unique_years) > 1:
                            return f"{unique_years[0]} - {unique_years[-1]}"

                    elif directive == "%d":  # Day of month
                        unique_days = dates.dt.strftime("%d").unique()
                        if len(unique_days) == 1:
                            return f"Day {unique_days[0]}"
                        elif len(unique_days) > 1:
                            return f"Day {unique_days[0]} - {unique_days[-1]}"

                    elif directive == "%j":  # Day of year
                        unique_days = dates.dt.strftime("%j").unique()
                        if len(unique_days) == 1:
                            return f"Day {int(unique_days[0])}"
                        elif len(unique_days) > 1:
                            return f"Day {int(unique_days[0])} - {int(unique_days[-1])}"

                    elif directive == "%U" or directive == "%W":  # Week number
                        unique_weeks = dates.dt.strftime("%U").unique()
                        if len(unique_weeks) == 1:
                            return f"Week {int(unique_weeks[0])}"
                        elif len(unique_weeks) > 1:
                            return f"Week {int(unique_weeks[0])} - {int(unique_weeks[-1])}"

                    else:
                        # Default: try to format with the directive
                        formatted = dates.dt.strftime(directive).unique()
                        if len(formatted) == 1:
                            return formatted[0]
                        elif len(formatted) > 1:
                            return f"{formatted[0]} - {formatted[-1]}"

            except Exception as e:
                logging.info(f"Error extracting temporal value from segment_start: {e}")

        # Fallback: Use directive mapping
        directive_mapping = {
            "%B": "Monthly",
            "%b": "Monthly",
            "%A": "Day of Week",
            "%a": "Day of Week",
            "%Y": "Yearly",
            "%d": "Daily",
            "%j": "Day of Year",
            "%U": "Weekly",
            "%W": "Weekly",
        }

        if directive in directive_mapping:
            return directive_mapping[directive]

    # Final fallback
    return "Temporal"


def _format_temporal_value(value: Any, frequency: str | None) -> str:
    """
    Format temporal value based on frequency type.
    Converts month numbers to names, day numbers to day names, etc.
    """
    # Month name mapping
    month_names = {
        1: "January",
        2: "February",
        3: "March",
        4: "April",
        5: "May",
        6: "June",
        7: "July",
        8: "August",
        9: "September",
        10: "October",
        11: "November",
        12: "December",
    }

    # Day of week mapping (0=Monday, 6=Sunday)
    day_names = {0: "Monday", 1: "Tuesday", 2: "Wednesday", 3: "Thursday", 4: "Friday", 5: "Saturday", 6: "Sunday"}

    # Handle month frequencies
    if frequency in ["M", "MS", "ME"]:
        # Check if value is a datetime-like object
        if hasattr(value, "month"):
            month_num = value.month
            year = value.year
            month_name = month_names.get(month_num, str(month_num))
            return f"{month_name} {year}"
        # Check if value is an integer (month number)
        elif isinstance(value, (int, float)) and 1 <= value <= 12:
            return month_names.get(int(value), str(value))
        # Check if value is a pandas Period
        elif hasattr(value, "strftime"):
            try:
                return value.strftime("%B %Y")
            except Exception as e:
                logging.info(f"{e}")
                pass

    # Handle day of week
    if frequency == "DOW":
        if isinstance(value, (int, float)) and 0 <= value <= 6:
            return day_names.get(int(value), str(value))
        elif isinstance(value, str):
            # Value might already be day name
            return value
        elif hasattr(value, "day_name"):
            return value.day_name()

    # Handle day of year
    if frequency == "DOY":
        if isinstance(value, (int, float)):
            return f"Day {int(value)}"

    # Handle daily frequency
    if frequency == "D":
        if hasattr(value, "strftime"):
            return value.strftime("%d %B %Y")

    # Handle yearly frequency
    if frequency in ["Y", "YS", "YE"]:
        if hasattr(value, "year"):
            return str(value.year)
        elif isinstance(value, (int, float)):
            return str(int(value))

    # Handle quarterly frequency
    if frequency == "Q":
        if hasattr(value, "quarter"):
            return f"Q{value.quarter} {value.year}"
        elif hasattr(value, "strftime"):
            month = value.month
            quarter = (month - 1) // 3 + 1
            return f"Q{quarter} {value.year}"

    # Handle weekly frequency
    if frequency == "W":
        if hasattr(value, "strftime"):
            return value.strftime("Week %U, %Y")

    # Default: convert to string
    return str(value)


def prepare_context_for_template(
    context: Any,
    template: DocxTemplate,
    box_h_cm: float = 6.5,
    box_w_cm: float = 11.11,
) -> dict:
    result = {}

    for key, value in context.items():
        if value is None:
            result[key] = value
            continue

        # Handle image fields - check if it's a valid image path
        if isinstance(value, str):
            normalized_path = remove_file_scheme(value)
            if os.path.exists(normalized_path):
                path_obj = Path(normalized_path)
                # If it's an image file, convert to InlineImage
                if path_obj.suffix.lower() in {".png", ".jpg", ".jpeg"}:
                    result[key] = InlineImage(template, normalized_path, width=Cm(box_w_cm), height=Cm(box_h_cm))
                else:
                    result[key] = value
            else:
                result[key] = value
        # Handle numeric fields - convert to integers
        elif isinstance(value, (int, float)):
            result[key] = int(value)
        else:
            result[key] = value

    return result


@task
def create_grouper_page(
    template_path: str,
    output_dir: str,
    context: dict[str, Any],
    filename: Optional[str] = None,
    validate_images: bool = True,
    box_h_cm: float = 6.5,
    box_w_cm: float = 11.11,
) -> str:
    template_path = remove_file_scheme(template_path)
    output_dir = remove_file_scheme(output_dir)

    # Validate paths
    if not template_path.strip():
        raise ValueError("template_path is empty after normalization")
    if not output_dir.strip():
        raise ValueError("output_directory is empty after normalization")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    if not filename:
        filename = f"context_{uuid.uuid4().hex}.docx"
    output_path = Path(output_dir) / filename

    if validate_images:
        for field_name, value in context.items():
            if isinstance(value, str) and Path(value).suffix.lower() in (".png", ".jpg", ".jpeg"):
                validate_image_path(field_name, value)

    try:
        tpl = DocxTemplate(template_path)
    except Exception as e:
        raise ValueError(f"Failed to load template: {e}")

    rendered_context = prepare_context_for_template(
        context=context,
        template=tpl,
        box_h_cm=box_h_cm,
        box_w_cm=box_w_cm,
    )
    try:
        tpl.render(rendered_context)
        tpl.save(output_path)
        return str(output_path)
    except Exception as e:
        raise ValueError(f"Failed to render or save document: {e}")


# merge word docs
@task
def merge_mapbook_files(
    cover_page_path: Annotated[str, Field(description="Path to the cover page .docx file")],
    context_page_items: Annotated[list[Any], Field(description="List of context page document paths to merge.")],
    output_dir: Annotated[str, Field(description="Directory where combined docx will be written")],
    filename: Annotated[Optional[str], Field(description="Optional output filename")] = None,
) -> Annotated[str, Field(description="Path to the combined .docx file")]:
    """
    Combine cover + context pages into a single DOCX.
    Orders context pages by calendar month if month names are detected.
    """
    import os
    import calendar
    from pathlib import Path
    from docx import Document
    from docxcompose.composer import Composer

    # Build month lookup from calendar module
    MONTH_LOOKUP = {name.lower(): idx for idx, name in enumerate(calendar.month_name) if name}
    MONTH_ABBR_LOOKUP = {name.lower(): idx for idx, name in enumerate(calendar.month_abbr) if name}

    def is_skip_sentinel(obj):
        """Check if object is a SkipSentinel."""
        return hasattr(obj, "__class__") and "SkipSentinel" in obj.__class__.__name__

    def extract_path(item):
        """Extract a valid file path from various item formats."""
        # Check for SkipSentinel
        if is_skip_sentinel(item):
            return None

        if isinstance(item, str):
            return item

        if isinstance(item, (list, tuple)):
            # Check for SkipSentinel in tuples
            for x in item:
                if is_skip_sentinel(x):
                    return None
                if isinstance(x, str) and os.path.exists(x):
                    return x
            # Fallback: return first string found
            for x in reversed(item):
                if isinstance(x, str):
                    return x
            return None

        return None

    def detect_month(path: str):
        """Detect month name or abbreviation in filename/path."""
        name = os.path.basename(path).lower()
        for month, idx in MONTH_LOOKUP.items():
            if month in name:
                return idx
        for abbr, idx in MONTH_ABBR_LOOKUP.items():
            if abbr in name:
                return idx
        return None

    # Normalize paths - filter out None and SkipSentinel
    normalized_paths = []
    logging.info(f"Context page items: {context_page_items}")
    for idx, item in enumerate(context_page_items):
        if item is None or is_skip_sentinel(item):
            logging.info(f"Skipping item {idx}: None or SkipSentinel")
            continue

        path = extract_path(item)
        if path is not None:
            normalized_paths.append(path)
        else:
            logging.error(f"Skipping item {idx}: Could not extract valid path")

    if not normalized_paths:
        logging.info("Warning: No valid context pages to merge, returning cover page only")
        # Just save cover page as output
        output_dir = remove_file_scheme(output_dir)
        os.makedirs(output_dir, exist_ok=True)
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"overall_report_{timestamp}.docx"
        output_path = Path(output_dir) / filename

        master = Document(cover_page_path)
        master.save(str(output_path))
        return str(output_path)

    if not os.path.exists(cover_page_path):
        raise FileNotFoundError(f"Cover page file not found: {cover_page_path}")

    for p in normalized_paths:
        if not os.path.exists(p):
            raise FileNotFoundError(f"Context page file not found: {p}")

    # Calendar-aware ordering
    with_month = []
    without_month = []
    for i, path in enumerate(normalized_paths):
        month_idx = detect_month(path)
        if month_idx is not None:
            with_month.append((month_idx, path))
        else:
            without_month.append((i, path))

    with_month.sort(key=lambda x: x[0])
    ordered_paths = [p for _, p in with_month] + [p for _, p in without_month]

    output_dir = remove_file_scheme(output_dir)
    if not output_dir.strip():
        raise ValueError("output_dir is empty after normalization")
    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"overall_report_{timestamp}.docx"

    output_path = Path(output_dir) / filename

    master = Document(cover_page_path)
    composer = Composer(master)

    for doc_path in ordered_paths:
        doc = Document(doc_path)
        composer.append(doc)

    composer.save(str(output_path))
    return str(output_path)
